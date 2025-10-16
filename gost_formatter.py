from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FORMAT_RULES = {
    'font': 'Times New Roman',
    'size': 14,
    'line_spacing': 1.5,
    'margin_left': 3.0,
    'margin_other': 2.0,
    'first_line_indent': 1.25
}

def add_page_number(doc):
    """Добавляет номера страниц с 3-й страницы"""
    sections = doc.sections
    for section in sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE \\* MERGEFORMAT"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)


def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FORMAT_RULES['font']
    font.size = Pt(FORMAT_RULES['size'])
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = FORMAT_RULES['line_spacing']
    paragraph_format.first_line_indent = Cm(FORMAT_RULES['first_line_indent'])


def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(FORMAT_RULES['margin_left'])
        section.right_margin = Cm(FORMAT_RULES['margin_other'])
        section.top_margin = Cm(FORMAT_RULES['margin_other'])
        section.bottom_margin = Cm(FORMAT_RULES['margin_other'])


def create_title_page(doc, topic: str):
    doc.add_paragraph().add_run("\n\n\n")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("НАЗВАНИЕ ВУЗА\n").bold = True
    title.add_run("Кафедра информационных технологий\n\n").bold = True
    title.add_run("ДОКЛАД\n").bold = True
    title.add_run(f'по теме: "{topic}"\n\n\n').bold = True
    title.add_run("Выполнил: студент гр. XYZ\n")
    title.add_run("Проверил: доцент Иванов И.И.\n\n\n")
    title.add_run("Москва 2024").bold = True
    doc.add_page_break()


def add_table_of_contents(doc):
    doc.add_paragraph("СОДЕРЖАНИЕ", style='Heading 1')
    toc_entries = [
        ("Введение", 3),
        ("1. Диагностика заболеваний с помощью ИИ", 4),
        ("2. Роботизированная хирургия", 6),
        ("3. Персонализированное лечение", 8),
        ("Заключение", 9),
        ("Список использованных источников", 10)
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        p.add_run(title).bold = False
        tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Cm(14))
        p.add_run(f"\t{page}")
    doc.add_page_break()


def add_section(doc, title: str, content: str):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(content)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_references(doc, topic: str):
    doc.add_heading("Список использованных источников", level=1)
    doc.add_paragraph("[1] Wikipedia. Artificial Intelligence in Medicine. URL: https://en.wikipedia.org/wiki/... (дата обращения: 2024)")
    doc.add_paragraph("[2] Сгенерировано ИИ-ассистентом на основе открытых данных.")


def save_document(doc, filename: str):
    doc.save(filename)
    return os.path.abspath(filename)
