from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FORMAT_RULES = {
    'font': 'Times New Roman',
    'size': 14,
    'line_spacing': 1.5,
    'margin_left': 3.0,
    'margin_other': 2.0,
    'first_line_indent': 1.25
}


def add_page_number(doc):
    """Добавляет номера страниц с 3-й страницы"""
    sections = doc.sections
    for section in sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE \\* MERGEFORMAT"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)


def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = FORMAT_RULES['font']
    font.size = Pt(FORMAT_RULES['size'])
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = FORMAT_RULES['line_spacing']
    paragraph_format.first_line_indent = Cm(FORMAT_RULES['first_line_indent'])


def set_margins(doc):
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(FORMAT_RULES['margin_left'])
        section.right_margin = Cm(FORMAT_RULES['margin_other'])
        section.top_margin = Cm(FORMAT_RULES['margin_other'])
        section.bottom_margin = Cm(FORMAT_RULES['margin_other'])


def create_title_page(doc, topic: str):
    # Очистка стиля: без отступов и выравнивание по центру
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)
    
    # Добавляем текст построчно, как в образце
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("АВТОНОМНАЯ НЕКОММЕРЧЕСКАЯ ОРГАНИЗАЦИЯ ПРОФЕССИОНАЛЬНАЯ ОБРАЗОВАТЕЛЬНАЯ ОРГАНИЗАЦИЯ\n").bold = True
    p.add_run("МОСКОВСКИЙ МЕЖДУНАРОДНЫЙ КОЛЛЕДЖ ЦИФРОВЫХ ТЕХНОЛОГИЙ\n").bold = True
    p.add_run("«АКАДЕМИЯ ТОП»\n").bold = True

    doc.add_paragraph()  # пустая строка

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ОТЧЁТ\n").bold = True

    doc.add_paragraph()  # пустая строка

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Уровень профессионального образования:\n")
    p.add_run("Среднее профессиональное образование\n\n")
    p.add_run("Программа подготовки специалистов среднего звена по специальности:\n")
    p.add_run("   09.02.07 Информационные системы и программирование\n\n")
    p.add_run("Квалификация:  Программист\n")
    p.add_run("Учебный предмет:  _______________________\n")
    p.add_run(f"Тема:  {topic}\n")
    p.add_run("наименование темы\n")

    doc.add_paragraph()  # пустая строка

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Преподаватель:\n")
    p.add_run("__________________        _________________\n")
    p.add_run("                         подпись             инициалы фамилия\n")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Обучающийся:\n")
    p.add_run("__________________        _________________\n")
    p.add_run("                         подпись дата        инициалы фамилия\n")

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Москва, 2025")


def add_table_of_contents(doc):
    doc.add_paragraph("СОДЕРЖАНИЕ", style='Heading 1')
    toc_entries = [
        ("Введение", 3),
        ("1. Диагностика заболеваний с помощью ИИ", 4),
        ("2. Роботизированная хирургия", 6),
        ("3. Персонализированное лечение", 8),
        ("Заключение", 9),
        ("Список использованных источников", 10)
    ]
    for title, page in toc_entries:
        p = doc.add_paragraph()
        p.add_run(title).bold = False
        tab_stop = p.paragraph_format.tab_stops.add_tab_stop(Cm(14))
        p.add_run(f"\t{page}")
    doc.add_page_break()


def add_section(doc, title: str, content: str):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(content)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_references(doc, topic: str):
    doc.add_heading("Список использованных источников", level=1)
    doc.add_paragraph("[1] Wikipedia. Artificial Intelligence in Medicine. URL: https://en.wikipedia.org/wiki/... (дата обращения: 2024)")
    doc.add_paragraph("[2] Сгенерировано ИИ-ассистентом на основе открытых данных.")


def save_document(doc, filename: str):
    doc.save(filename)
    return os.path.abspath(filename)
